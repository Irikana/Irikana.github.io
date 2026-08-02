# -*- coding: utf-8 -*-
"""
生成规范格式的 Word 论文文档
专业文献写作课程作业
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FIG_DIR = r'g:\PClite\shepherdsLibrary\Irikana.github.io\docs\paper-figures'
OUT_PATH = r'g:\PClite\shepherdsLibrary\Irikana.github.io\docs\运动在对称操作下的不变性_韦仁杰.docx'

doc = Document()

# ============================================================
# 全局样式设置（仿学术期刊格式）
# ============================================================
# 页面设置：A4，标准页边距
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 默认字体：宋体小四，行距 1.5
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)

def set_run_font(run, cn='宋体', en='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(text, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    set_run_font(r, cn='黑体', en='Times New Roman', size=size, bold=True)
    return p

def add_subtitle(text, size=13, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_run_font(r, cn='楷体', en='Times New Roman', size=size, italic=True)
    return p

def add_h2(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, cn='黑体', en='Times New Roman', size=size, bold=True)
    return p

def add_h3(text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, cn='黑体', en='Times New Roman', size=size, bold=True)
    return p

def add_body(text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2字符
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=12)
    return p

def add_body_mixed(parts, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """parts: list of (text, dict) where dict may contain 'italic','bold','sub','sup'"""
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    for text, fmt in parts:
        r = p.add_run(text)
        set_run_font(r, size=12,
                     italic=fmt.get('italic', False),
                     bold=fmt.get('bold', False))
        if fmt.get('sub'):
            r.font.subscript = True
        if fmt.get('sup'):
            r.font.superscript = True
    return p

def add_formula(text, number=None):
    """居中公式段落，text 用斜体"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, en='Times New Roman', size=12, italic=True)
    if number:
        # 公式编号右对齐
        r2 = p.add_run('\t\t(' + number + ')')
        set_run_font(r2, size=12)
    return p

def add_figure(img_path, caption, width_cm=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run()
    r.add_picture(img_path, width=Cm(width_cm))
    # 图题
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(caption)
    set_run_font(r2, cn='宋体', en='Times New Roman', size=10.5, bold=True)

def add_ref(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    r = p.add_run(text)
    set_run_font(r, size=10.5)

# ============================================================
# 标题区
# ============================================================
add_title('运动在对称操作下的不变性', size=18, space_after=4)
add_subtitle('——一种基于对称操作的问题转化方法', size=13)

# 作者信息
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('韦仁杰')
set_run_font(r, cn='楷体', size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('（物理学院 2024 级本科生）')
set_run_font(r, cn='楷体', size=10.5)

# ============================================================
# 摘要
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.right_indent = Cm(0.5)
r1 = p.add_run('摘  要：')
set_run_font(r1, cn='黑体', size=10.5, bold=True)
r2 = p.add_run('本文提出并阐释一个关于对称操作的猜想：一个问题可以通过有限的对称操作转化为另一个问题，给出完全相同的结果，但以完全不同的方式呈现。这一猜想将对称性从"系统自身的性质"拓展为"问题之间转化的工具"。通过两个案例展示该方法：其一，匀速游动的鱼所受水的作用力，通过惯性系切换（伽利略变换）等价转化为静止鱼所受的力；其二，凹凸轨道上两小球谁先到达终点的问题，通过旋转装置构造对称性、再分析对称性破缺，无需详细计算即可定性得出结论。在此基础上，进一步提出"简单模型拼凑"的方法构想——通过将简单的对称刚体或对称场逐渐对称地加入到系统中，直至与目标问题一致，从而以更少的工作量得到答案。本文亦回顾了对称性在物理学史中的发展脉络，以说明本工作所处的思想背景。')
set_run_font(r2, cn='宋体', size=10.5)

# 关键词
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(12)
r1 = p.add_run('关键词：')
set_run_font(r1, cn='黑体', size=10.5, bold=True)
r2 = p.add_run('对称操作；问题转化；伽利略变换；对称性破缺；简单模型')
set_run_font(r2, cn='宋体', size=10.5)

# ============================================================
# 1. 引言
# ============================================================
add_h2('1  引言')

add_body('对称性的故事，始于人类对宇宙和谐最古老的直觉。毕达哥拉斯学派说"万物皆数"，柏拉图用五种正多面体对应五种元素——在他们的眼中，对称不是数学工具，而是宇宙的本质。开普勒是一个极其虔诚的毕达哥拉斯-柏拉图主义者，他花了一辈子寻找行星轨道之间的完美几何对称：25岁时，他用五种柏拉图立体嵌套球壳来"解释"六颗行星的轨道间距，认为这是上帝创造宇宙的蓝图；在《世界的和谐》中，他为每颗行星谱写旋律，将角速度之比对应为音乐的协和音程。讽刺的是，他毕生追寻的"完美对称"最终被证明是错的——行星不是六个而是八个，轨道间距与正多面体嵌套并不精确吻合。但正是为了验证这些"和谐"假设，他不得不使用第谷·布拉赫的精确数据，而精确数据迫使他放弃了圆形轨道，发现了真正的行星运动定律。他追寻的幻象引导他找到了真实。')

add_body('1632年，伽利略在"船舱"思想实验中，第一次明确表述了一个物理学对称性：在所有匀速直线运动的参考系中，物理定律保持不变[1]。牛顿在《自然哲学的数学原理》中进一步阐明了这一点[2]。然而，此后的两百年里，没有人意识到"对称性"本身可以成为一个独立的物理概念——它只是运动方程的一个附带性质。')

add_body('转折发生在1915年。爱因斯坦发表广义相对论后，一个核心问题困扰着整个物理学界：能量守恒定律是否仍然成立？希尔伯特和克莱因将这个问题交给了哥廷根大学一位连正式教职都没有的数学家——艾米·诺特。诺特用了不到两年，不仅彻底解决了广义相对论的能量守恒问题，而且证明了一个更深刻的定理：每一个连续对称性，都对应一个守恒律[3]。时间平移对应能量守恒，空间平移对应动量守恒，空间旋转对应角动量守恒——这些原本各自独立的实验定律，原来共享同一个起源。')

add_body('诺特定理让物理学界意识到对称性的力量。1927年，维格纳提出宇称守恒定律[4]——物理定律在镜像变换下应保持不变。此后三十年，没有人质疑过它。但杨振宁和李政道在1956年检查后发现了一个惊人的事实：在弱相互作用领域，没有任何一个实验曾经检验过宇称是否守恒[5]。年底，吴健雄用极化钴-60的β衰变实验给出了答案：宇宙并不是完美对称的[6]。费曼后来反思道：物理学家最大的错误之一，就是将美学偏好当作物理定律来对待。')

add_body('但对称性的故事并未因此结束，反而走向了更深刻的方向。1957年，BCS超导理论揭示了一种奇特的现象：描述系统的拉格朗日量是对称的，但基态破坏了这种对称——自发对称破缺。南部阳一郎将这一思想从凝聚态物理推广到粒子物理[7]。随后，戈德斯通定理预言了无质量玻色子[8]，最终在1964年，三组物理学家独立提出了希格斯机制[9]。1967年，温伯格将希格斯机制与杨-米尔斯理论结合，建立了电弱统一理论[10]。宇宙的秩序来源于对称性，宇宙的多样性来源于对称破缺。而这一切，始于1918年诺特写下的那十几页纸。')

add_body('本文的工作正是在这一思想背景下展开的。这是一个已经进行了很久、并且到现在仍在进行的工作。我们试图将对称性不仅理解为系统自身的性质，更理解为一种可以在问题之间进行转化的操作工具。下面，我们先陈述核心猜想，再以两个案例加以展示，最后讨论这一方法的普遍意义与局限。')

# ============================================================
# 2. 核心猜想
# ============================================================
add_h2('2  核心猜想')

# 猜想框
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.left_indent = Cm(1.5)
p.paragraph_format.right_indent = Cm(1.5)
r = p.add_run('猜想：一个问题可以通过有限的对称操作转化为另一个问题，给出完全相同的结果，但以完全不同的方式呈现。')
set_run_font(r, cn='楷体', size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.left_indent = Cm(1.5)
p.paragraph_format.right_indent = Cm(1.5)
r = p.add_run('"One problem could be converted to another problem by taking limited symmetrical operation, which would give me a result absolutely the same but could appear by totally different ways."')
set_run_font(r, en='Times New Roman', size=10.5, italic=True)

add_body('我们最开始直接使用了"有限"一词，是直觉使然。后来才想明白，强调有限仅仅是因为无限是没有意义的——你没办法做出无限多的操作。这一猜想的核心在于：对称操作不改变问题的物理实质，只改变问题的呈现方式。因此，通过选择合适的对称操作，我们可以把一个困难的问题转化为一个等价但更简单的问题。')

add_body('需要说明的是，这里的"对称操作"取较宽的含义：既包括经典的几何对称操作（旋转、反射、平移），也包括参考系变换（伽利略变换），以及通过人为构造装置来建立或破缺对称性的操作。下面两个案例分别展示了不同类型的对称操作如何实现问题的转化。')

# ============================================================
# 3. 案例一
# ============================================================
add_h2('3  案例一：匀速游动的鱼')

add_h3('3.1  问题')

add_body('一条鱼以恒定速度在水中向左匀速直线游动。问：水对鱼的作用力方向如何？')

add_figure(os.path.join(FIG_DIR, 'fig1_fish_moving.png'),
           '图1  匀速向左游动的鱼', width_cm=13)

add_body('面对这个问题，若直接从流体力学角度分析鱼的受力，需要考虑鱼游动时尾鳍的推进力、水的阻力、鱼身体的形状阻力等诸多因素，问题相当复杂。然而，借助一次对称操作，问题可以瞬间简化。')

add_h3('3.2  对称操作：惯性系切换')

add_body('牛顿第一定律的深层含义在于：匀速直线运动与静止在物理上完全等价——无法从实验上区分二者。这本身就是一条关于对称性的陈述：所有惯性参考系在物理上是等价的，物理定律在惯性系之间的伽利略变换下保持不变[1,11]。')

add_body('设鱼在水面参考系（实验室系）中以速度 v 向左运动。伽利略变换给出两惯性系之间的坐标变换关系：')

add_formula("x' = x − vt,    t' = t", "(1)")

add_body('在此变换下，加速度不变（a′ = a），从而牛顿第二定律 F = ma 的形式不变。我们将参考系切换到与鱼同速运动的惯性系。在这一参考系中，鱼静止于水中，而水以相同的速率向右流过鱼的身体。于是原来的问题"匀速向左游动的鱼受力如何"被等价转化为"在流动水中静止的鱼受力如何"。')

add_figure(os.path.join(FIG_DIR, 'fig2_fish_rest_frame.png'),
           '图2  切换到鱼静止的参考系（伽利略变换）', width_cm=13)

add_body('在新的问题中，鱼静止，根据平衡条件，鱼所受合力为零：')

add_formula("ΣF = 0    ⇒    F水 + mg = 0    ⇒    F水 = −mg", "(2)")

add_body('鱼受到向下的重力 mg，因此水对鱼的作用力必须竖直向上，以平衡重力。答案立刻可得：水对鱼的作用力方向竖直向上。')

add_h3('3.3  分析')

add_body('这一案例中，对称操作是伽利略变换（惯性系切换）。问题的物理实质——鱼在水中匀速运动时的受力——并未改变，但问题的呈现方式从"运动中的鱼"变成了"静止中的鱼"。后者由于可以直接使用平衡条件，远比前者简单。这正是核心猜想的体现：通过一次有限的对称操作，问题被转化为一个等价但更易处理的形式，给出了完全相同的结果。')

add_body('值得注意的是，这里所用的对称性——伽利略不变性——正是物理学中最古老、最基础的对称性之一。它早在1632年就由伽利略在船舱思想实验中表述[1]。本案例的意义在于：它不仅是对称性的一个应用，更展示了"对称操作作为问题转化工具"这一思想的最纯粹形式。')

# ============================================================
# 4. 案例二
# ============================================================
add_h2('4  案例二：凹凸轨道')

add_h3('4.1  问题')

add_body('有两个相同的圆弧轨道，一凹一凸，关于一条水平线对称。两个相同的小球以相同的初速度 v₀ 分别从两个轨道的左端出发，沿轨道无摩擦地运动到右端。问：哪个小球先到达终点？')

add_figure(os.path.join(FIG_DIR, 'fig3_tracks_problem.png'),
           '图3  凹凸轨道问题', width_cm=13)

add_body('这是一个看似需要详细计算的问题。若直接用能量守恒和运动学方程求解每个小球到达终点的时间，需要积分计算，过程繁琐。然而，借助对称操作的构造与破缺分析，可以定性得出结论而几乎无需计算。')

add_body('作为对比，我们先写出标准的定量方法。对任一小球，由机械能守恒：')

add_formula("½mv₀² + mgh(x) = ½mv(x)²", "(3)")

add_body('其中 h(x) 是小球在轨道上位置 x 处的高度（以终点为参考）。由此得速率：')

add_formula("v(x) = √(v₀² + 2gh(x))", "(4)")

add_body('到达终点的时间为 T = ∫ ds/v(x)，其中 ds 为轨道弧长微元。对凹凸两轨道，h(x) 的符号相反，因此 v(x) 不同，T 也不同。这一积分通常没有简单的解析形式，需要数值计算。下面我们展示如何用对称操作避开这一计算。')

add_h3('4.2  对称操作的构造')

add_body('关键观察在于：两个轨道本身关于水平线是对称的，但重力方向竖直向下，并不沿这条对称轴方向——正是重力破坏了系统原有的对称性。如果我们能够构造一种情形，使得重力方向与对称轴共线，那么系统将恢复完全对称，两球将同时到达。')

add_body('为此，我们进行如下操作：')

add_body('第一步：将整个装置（两个轨道连同小球）旋转90°，使原来的水平对称轴变为竖直方向。此时，重力方向（竖直向下）与对称轴共线。', indent=False)

add_body('第二步：在旋转后的装置中，将凹轨道的实线向左平移，在其右侧添加一条凹轨道的虚线镜像，使之与凸轨道实线对称；凸轨道侧同理添加虚线镜像。如此一来，两组轨道形成了一个"双壁"通道，整体构型关于竖直对称轴完全对称。', indent=False)

add_figure(os.path.join(FIG_DIR, 'fig4_symmetric_config.png'),
           '图4  旋转90°并添加镜像轨道后，系统关于竖直轴完全对称', width_cm=10)

add_body('第三步：在这一完全对称的构型中，两个小球的运动是完全对称的，因此它们同时到达终点。这是一个无需计算即可得出的结论。', indent=False)

add_h3('4.3  对称性破缺分析')

add_body('第四步：现在，将装置旋转回原来的位置。这一旋转操作等价于让重力方向偏离了对称轴——也就是说，我们破缺了第三步中建立的对称性。', indent=False)

add_body('第五步：对称性破缺后，重力场偏向了凹轨道一侧。对于凹轨道上的小球，重力在运动方向上的分量做更多的正功（小球先下坡，重力势能更多转化为动能），因此其平均速度更大；对于凸轨道上的小球，重力先做负功（先上坡），平均速度较小。因此，凹轨道上的小球先到达终点。', indent=False)

add_figure(os.path.join(FIG_DIR, 'fig5_symmetry_breaking.png'),
           '图5  旋转回原位即对称性破缺，凹轨道小球先到达', width_cm=13)

add_h3('4.4  分析')

add_body('这一案例展示了比案例一更复杂的对称操作：它不是简单地利用已有的对称性，而是主动地构造对称性（通过旋转和添加镜像装置），再通过破缺对称性来定性分析。整个推理过程可以概括为一条原则：')

# 原则框
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.left_indent = Cm(1.5)
p.paragraph_format.right_indent = Cm(1.5)
r = p.add_run('对称 → 对称结果；破缺对称性 → 偏离对称结果；偏离方向 = 破缺方向。')
set_run_font(r, cn='楷体', size=12, bold=True)

add_body('这一方法的威力在于：它将一个需要定量计算的问题，转化为一个只需定性分析的对称性论证。我们不需要写出运动方程，不需要做积分，只需要识别对称性、构造对称性、再分析对称性如何被破缺。这正符合核心猜想的精神——通过有限的对称操作，问题被转化为另一种呈现方式，以完全不同的途径得到了相同的结果。')

# ============================================================
# 5. 简单模型方法
# ============================================================
add_h2('5  "简单模型"方法')

add_body('在上述两个案例的基础上，我们进一步提出一种更一般的方法构想，称之为"简单模型拼凑"。所谓"拼凑"，是指将不同物理系统中的简单模型，通过有限的对称操作进行组合与映射，使得一个系统的已知结论可以直接迁移到另一个系统，而无需重新推导。')

add_body('"简单模型"是那些最基本的、简单的对称刚体或对称的场，或者是基本粒子（我们对这部分物理是陌生的，所以只是猜测）。例如，当我们在研究自由落体的物体如何运动时，可以往任意的方向上添加一个和引力场相同性质的场——这会改变问题，但不会改变运动的形式，物体会换一个方向继续做落体运动。我们认为从这个角度思考问题是有价值的，这是我们做这些工作的原因。')

add_body('以自由落体为例。设物体在均匀引力场 g 中做自由落体运动，加速度为 g，方向向下。若我们添加一个与 g 同性质的均匀场 g′，方向任意，则物体所受总加速度为 a = g + g′。运动的形式仍是匀加速直线运动，只是方向和大小发生了改变。用方程表示：')

add_formula("r(t) = r₀ + v₀t + ½(g + g′)t²", "(5)")

add_body('这一形式与原自由落体方程 r(t) = r₀ + v₀t + ½gt² 在结构上完全一致，只是加速度矢量被替换。这意味着：关于自由落体的所有已知结论（如抛体运动的轨迹方程、飞行时间、射程公式等）都可以通过对称地"添加场"这一操作，直接迁移到新的问题中，而无需重新求解。')

add_body('这一方法的核心思想是：通过将简单模型逐渐对称地加入到我们的系统当中，直到系统与目标问题类似或完全一致，从而能够以更少的工作量得到答案。简单模型本身是高度对称的、易于分析的；通过对称操作将它们组合或映射到复杂系统中，复杂系统的某些性质就可以从简单模型的已知结论中直接读出，而不必从头求解。')

add_body('这一构想目前仍处于猜想阶段，尚未形成严格的理论框架。但它指出了一个值得探索的方向：对称操作不仅可以用于在问题之间转化（如案例一、二所示），还可以用于从简单到复杂的构造。如果这一方法能够系统化，它可能为物理问题的求解提供一种补充性的思路——不依赖于直接的数学计算，而是依赖于对称性的构造与映射。')

# ============================================================
# 6. 方法流程
# ============================================================
add_h2('6  方法流程图')

add_body('综合上述两个案例与"简单模型"方法，可以将本工作提出的问题转化方法归纳为如下流程：')

add_figure(os.path.join(FIG_DIR, 'fig6_method_flowchart.png'),
           '图6  对称操作问题转化方法流程', width_cm=14)

add_body('如图6所示，原问题首先经过对称性的识别或构造，然后通过有限步对称操作转化为一个等价但更简单的新问题，从而得到相同的结果。当对称性被破缺时，可以通过分析破缺的方向，定性推断实际结果相对于对称情形的偏离。')

# ============================================================
# 7. 讨论
# ============================================================
add_h2('7  讨论')

add_h3('7.1  与诺特定理的关系')

add_body('诺特定理建立的是"对称性—守恒律"的对应：每一个连续对称性对应一个守恒律[3]。本工作所讨论的并非守恒律，而是"对称性—问题转化"的对应：一个对称操作可以将一个问题转化为另一个等价的问题。两者都属于对称性的应用，但着眼点不同。诺特定理关注的是单一系统内部对称性所蕴含的不变量；本工作关注的是不同问题之间通过对称操作建立的等价关系。这两种视角并不矛盾，而是互补的。')

add_h3('7.2  与对称性破缺的关系')

add_body('案例二直接运用了对称性破缺的思想。在物理学史上，对称性破缺从凝聚态物理（BCS超导）出发，经南部阳一郎推广到粒子物理[7]，最终发展为希格斯机制[9]和电弱统一理论[10]。本工作中，对称性破缺被用作一种定性分析工具：先建立完全对称的情形（此时结论显然），再分析对称性如何被破缺以及破缺的方向，从而推断实际情形相对于对称情形的偏离。这一思路与物理学史上"从对称到破缺"的思想一脉相承，但应用层面不同——前者用于解释基本粒子的质量起源，后者用于简化经典力学问题的求解。')

add_h3('7.3  "有限"的含义')

add_body('猜想中强调"有限的对称操作"，这一点值得说明。我们最初使用"有限"一词是出于直觉，后来才理解其必要性：无限多的对称操作在物理上是没有意义的，因为你没办法做出无限多的操作。更深一层地说，如果一个问题的转化需要无限多步操作，那么这种转化实际上是不可执行的——它没有给出比直接求解原问题更简单的途径。因此，"有限"不是猜想的限制，而是猜想的内在要求：只有有限步的可执行操作，才构成真正有意义的问题转化。')

add_h3('7.4  方法的局限')

add_body('本工作目前存在明显的局限。第一，两个案例虽然展示了猜想的可行性，但尚未给出"什么样的对称操作适用于什么样的问题"的一般判据。第二，"简单模型拼凑"的方法目前仅停留在构想层面，缺乏严格的数学表述和可操作的算法。第三，本文的分析以定性为主，未给出定量的适用范围与误差估计。这些局限正是后续工作需要弥补的。')

# ============================================================
# 8. 结论
# ============================================================
add_h2('8  结论')

add_body('本文提出了一个关于对称操作的猜想：一个问题可以通过有限的对称操作转化为另一个问题，给出完全相同的结果，但以完全不同的方式呈现。我们通过两个案例展示了这一猜想：匀速游动的鱼所受的力，通过惯性系切换等价转化为静止鱼所受的力；凹凸轨道上两小球谁先到达的问题，通过旋转构造对称性、再分析对称性破缺，定性得出凹轨道先到的结论。在此基础上，我们进一步提出了"简单模型拼凑"的方法构想——通过对称操作将简单模型组合到复杂系统中，以减少求解的工作量。')

add_body('这一工作的意义不在于给出新的物理定律，而在于提供一种看待物理问题的方式：将对称性不仅视为系统自身的性质，更视为问题之间转化的工具。从伽利略的船舱到诺特的定理，从宇称不守恒到自发对称破缺，物理学史一再表明，对称性是理解自然最深刻的视角之一。本工作试图在这一传统中，探索对称性作为问题求解工具的可能性。这一探索仍在进行中。')

# ============================================================
# 参考文献
# ============================================================
add_h2('参考文献')

refs = [
    '[1] Galilei G. Dialogo sopra i due massimi sistemi del mondo: tolemaico e copernicano[M]. Firenze: Landini, 1632.',
    '[2] Newton I. Philosophiæ Naturalis Principia Mathematica[M]. Londini: Jussu Societatis Regiæ, 1687.',
    '[3] Noether E. Invariante Variationsprobleme[J]. Nachrichten von der Gesellschaft der Wissenschaften zu Göttingen, Mathematisch-Physikalische Klasse, 1918: 235–257.',
    '[4] Wigner E P. Einige Folgerungen aus der Schrödingerschen Theorie für die Termstrukturen[J]. Zeitschrift für Physik, 1927, 43: 624–652.',
    '[5] Lee T D, Yang C N. Question of Parity Conservation in Weak Interactions[J]. Physical Review, 1956, 104(1): 254–258.',
    '[6] Wu C S, Ambler E, Hayward R W, et al. Experimental Test of Parity Conservation in Beta Decay[J]. Physical Review, 1957, 105(4): 1413–1415.',
    '[7] Nambu Y. Quasi-Particles and Gauge Invariance in the Theory of Superconductivity[J]. Physical Review, 1960, 117(3): 648–663.',
    '[8] Goldstone J. Field Theories with Superconductor Solutions[J]. Il Nuovo Cimento, 1961, 19(1): 154–164.',
    '[9] Higgs P W. Broken Symmetries and the Masses of Gauge Bosons[J]. Physical Review Letters, 1964, 13(16): 508–509.',
    '[10] Weinberg S. A Model of Leptons[J]. Physical Review Letters, 1967, 19(21): 1264–1266.',
    '[11] Feynman R P, Leighton R B, Sands M. The Feynman Lectures on Physics: Vol. I[M]. Reading, MA: Addison-Wesley, 1963.',
    '[12] Goldstein H, Poole C, Safko J. Classical Mechanics[M]. 3rd ed. San Francisco: Addison-Wesley, 2002.',
    '[13] Landau L D, Lifshitz E M. Mechanics[M]. 3rd ed. Oxford: Butterworth-Heinemann, 1976.',
]
for r in refs:
    add_ref(r)

# 保存
doc.save(OUT_PATH)
print('Word 文档已生成：', OUT_PATH)
print('文件大小：', os.path.getsize(OUT_PATH), 'bytes')
